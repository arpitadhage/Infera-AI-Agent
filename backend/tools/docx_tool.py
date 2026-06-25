from docx import Document

def extract_text_from_docx(state):
    file_path = state["file_path"]

    doc = Document(file_path)

    text = ""

    for para in doc.paragraphs:
        text += para.text + "\n"

    state["contents"]["docx"] = [{
        "text": text.strip(),
        "pages": len(doc.paragraphs)
    }]

    return state